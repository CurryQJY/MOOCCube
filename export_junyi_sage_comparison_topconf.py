"""Export Junyi SAGE comparison tables in the compact top-conference style.

This mirrors the earlier MOOCCube SAGE comparison table:
  output/doc/final_narrow_topconf/maskfalse_sage_comparison_narrow_topconf.docx

Outputs:
  output/doc/junyi_results/junyi_sage_tail0p01_scope_true_comparison_narrow_topconf.docx
  output/doc/junyi_results/junyi_sage_tail0p01_scope_true_comparison_cold.csv
  output/doc/junyi_results/junyi_sage_tail0p01_scope_true_comparison_hot.csv
  output/doc/junyi_results/junyi_sage_course_scope_false_comparison_narrow_topconf.docx
  output/doc/junyi_results/junyi_sage_course_scope_false_comparison_cold.csv
  output/doc/junyi_results/junyi_sage_course_scope_false_comparison_hot.csv
"""

from __future__ import annotations

import csv
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"D:\DeskTop\MOOCCube")
OUT_DIR = ROOT / "output" / "doc" / "junyi_results"
SCOPE_TRUE_DOCX = OUT_DIR / "junyi_sage_tail0p01_scope_true_comparison_narrow_topconf.docx"
SCOPE_TRUE_COLD_CSV = OUT_DIR / "junyi_sage_tail0p01_scope_true_comparison_cold.csv"
SCOPE_TRUE_HOT_CSV = OUT_DIR / "junyi_sage_tail0p01_scope_true_comparison_hot.csv"
SCOPE_FALSE_DOCX = OUT_DIR / "junyi_sage_course_scope_false_comparison_narrow_topconf.docx"
SCOPE_FALSE_COLD_CSV = OUT_DIR / "junyi_sage_course_scope_false_comparison_cold.csv"
SCOPE_FALSE_HOT_CSV = OUT_DIR / "junyi_sage_course_scope_false_comparison_hot.csv"

OURS_DIR = (
    ROOT
    / "outputs"
    / "junyi"
    / "mask_ablation"
    / "mask_tt"
    / "strict_item_cold_balanced_thr1_seed_2025"
)
SAGE_SCOPE_TRUE_DIR = (
    ROOT
    / "outputs"
    / "junyi"
    / "sage_lite_v1"
    / "S1_tailratio_grid_seed2025"
    / "r0p01"
    / "strict_item_cold_balanced_thr1_seed_2025"
)
SAGE_SCOPE_FALSE_DIR = (
    ROOT
    / "outputs"
    / "junyi"
    / "sage_lite_v1"
    / "S2_course_scope_false_tail0p01_e60_seed2025"
    / "strict_item_cold_balanced_thr1_seed_2025"
)

COLD_METRICS = [
    ("Recall@5", "full_cold_item_macro_r5"),
    ("Recall@10", "full_cold_item_macro_r10"),
    ("Recall@20", "full_cold_item_macro_r20"),
    ("NDCG@5", "full_cold_item_macro_n5"),
    ("NDCG@10", "full_cold_item_macro_n10"),
    ("NDCG@20", "full_cold_item_macro_n20"),
]

HOT_METRICS = [
    ("Hot R@5", "full_hot_item_macro_r5"),
    ("Hot R@10", "full_hot_item_macro_r10"),
    ("Hot R@20", "full_hot_item_macro_r20"),
    ("Hot N@5", "full_hot_item_macro_n5"),
    ("Hot N@10", "full_hot_item_macro_n10"),
    ("Hot N@20", "full_hot_item_macro_n20"),
]


def read_final_row(run_dir: Path) -> dict[str, str]:
    path = run_dir / "final_fullrank_usim_feedback_fast3_content_delta_static.csv"
    if not path.exists():
        raise FileNotFoundError(path)
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return next(csv.DictReader(f))


def read_manifest(run_dir: Path) -> dict:
    path = run_dir / "static_protocol_manifest.json"
    if not path.exists():
        raise FileNotFoundError(path)
    return json.loads(path.read_text(encoding="utf-8"))


def require_config(sage_dir: Path, *, course_scope_false: bool) -> None:
    ours_cfg = read_manifest(OURS_DIR).get("model_config", {})
    sage_cfg = read_manifest(sage_dir).get("model_config", {})

    expected_ours = {
        "mask_known_pos_neg": True,
        "mask_same_item_neg": True,
        "use_sage_lite": False,
    }
    for key, value in expected_ours.items():
        if ours_cfg.get(key) is not value:
            raise ValueError(f"Ours config mismatch: {key}={ours_cfg.get(key)!r}, expected {value!r}")

    expected_sage = {
        "mask_known_pos_neg": True,
        "mask_same_item_neg": True,
        "use_sage_lite": True,
        "sage_only_cold_or_tail": True,
        "feedback_course_only_cold": not course_scope_false,
        "feedback_course_sample_only_cold": not course_scope_false,
        "prereq_aux_only_cold": not course_scope_false,
    }
    for key, value in expected_sage.items():
        if sage_cfg.get(key) is not value:
            raise ValueError(f"SAGE config mismatch: {key}={sage_cfg.get(key)!r}, expected {value!r}")
    if abs(float(sage_cfg.get("sage_tail_pop_ratio", -1.0)) - 0.01) > 1e-12:
        raise ValueError(f"SAGE ratio mismatch: {sage_cfg.get('sage_tail_pop_ratio')!r}, expected 0.01")


def fmt(value: float) -> str:
    return f"{value:.4f}"


def fmt_imp(value: float) -> str:
    return f"{value:+.2f}%"


def build_rows(metrics: list[tuple[str, str]], sage_dir: Path) -> list[dict[str, str]]:
    ours = read_final_row(OURS_DIR)
    sage = read_final_row(sage_dir)
    rows = []
    for model, row in (("Ours", ours), ("Ours + SAGE", sage)):
        out = {"Model": model}
        for label, key in metrics:
            out[label] = fmt(float(row[key]))
        rows.append(out)

    imp = {"Model": "Imp."}
    for label, key in metrics:
        base = float(ours[key])
        value = float(sage[key])
        imp[label] = fmt_imp(0.0 if abs(base) < 1e-12 else ((value / base) - 1.0) * 100.0)
    rows.append(imp)
    return rows


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, spec in kwargs.items():
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in spec.items():
            element.set(qn(f"w:{key}"), str(value))


def clear_cell_borders(cell) -> None:
    nil = {"val": "nil", "sz": "0", "space": "0", "color": "FFFFFF"}
    set_cell_border(cell, top=nil, bottom=nil, left=nil, right=nil, insideH=nil, insideV=nil)


def add_rule(cell, edge: str, width: str = "6") -> None:
    set_cell_border(cell, **{edge: {"val": "single", "sz": width, "space": "0", "color": "000000"}})


def set_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=32, start=24, bottom=32, end=24) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_text(cell, text: str, *, bold=False, italic=False, size=8.0, align=None) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def style_cell(cell) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    clear_cell_borders(cell)


def apply_rules(table) -> None:
    n_rows = len(table.rows)
    n_cols = len(table.columns)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            clear_cell_borders(cell)
            if col_idx < n_cols - 1:
                add_rule(cell, "right", "5")
            if row_idx == 0:
                add_rule(cell, "top", "9")
                add_rule(cell, "bottom", "7")
            if row_idx == n_rows - 1:
                add_rule(cell, "bottom", "9")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(9.2)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(7.5)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_comparison_table(
    doc: Document,
    rows: list[dict[str, str]],
    metrics: list[tuple[str, str]],
    *,
    caption: str,
    note: str,
) -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=len(rows) + 1, cols=1 + len(metrics))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.23] + [0.62] * len(metrics)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_width(cell, widths[idx])
            style_cell(cell)

    headers = ["Model"] + [label for label, _ in metrics]
    for idx, header in enumerate(headers):
        add_text(table.rows[0].cells[idx], header, bold=(idx == 0), italic=(idx > 0), size=7.8)

    for row_idx, row in enumerate(rows, start=1):
        add_text(table.rows[row_idx].cells[0], row["Model"], size=8.2)
        for col_idx, (label, _) in enumerate(metrics, start=1):
            add_text(table.rows[row_idx].cells[col_idx], row[label], size=8.0)

    apply_rules(table)
    add_note(doc, note)


def write_csv(rows: list[dict[str, str]], metrics: list[tuple[str, str]], path: Path) -> None:
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Model"] + [label for label, _ in metrics])
        for row in rows:
            writer.writerow([row["Model"]] + [row[label] for label, _ in metrics])


def build_docx(
    cold_rows: list[dict[str, str]],
    hot_rows: list[dict[str, str]],
    *,
    out_docx: Path,
    note: str,
) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(9)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_comparison_table(
        doc,
        cold_rows,
        COLD_METRICS,
        caption="Table S1: Junyi cold-item comparison on the balanced static item-cold split.",
        note=note,
    )
    add_comparison_table(
        doc,
        hot_rows,
        HOT_METRICS,
        caption="Table S2: Junyi hot-item comparison on the balanced static item-cold split.",
        note=note,
    )
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    return save_docx(doc, out_docx)


def save_docx(doc: Document, out_docx: Path) -> Path:
    try:
        doc.save(out_docx)
        return out_docx
    except PermissionError:
        fallback = out_docx.with_name(f"{out_docx.stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{out_docx.suffix}")
        doc.save(fallback)
        print(f"Locked: {out_docx}; wrote fallback {fallback}")
        return fallback


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    exports = [
        {
            "sage_dir": SAGE_SCOPE_TRUE_DIR,
            "course_scope_false": False,
            "docx": SCOPE_TRUE_DOCX,
            "cold_csv": SCOPE_TRUE_COLD_CSV,
            "hot_csv": SCOPE_TRUE_HOT_CSV,
            "note": (
                "All values are seed-2025 test results under full-ranking item-macro evaluation. "
                "Ours uses MaskKnownPosNeg=True and MaskSameItemNeg=True. "
                "Ours + SAGE uses tail ratio 0.01 with course feedback, course sampling, and prereq auxiliary loss restricted to cold items; "
                "Imp. reports the relative difference over Ours."
            ),
        },
        {
            "sage_dir": SAGE_SCOPE_FALSE_DIR,
            "course_scope_false": True,
            "docx": SCOPE_FALSE_DOCX,
            "cold_csv": SCOPE_FALSE_COLD_CSV,
            "hot_csv": SCOPE_FALSE_HOT_CSV,
            "note": (
                "All values are seed-2025 test results under full-ranking item-macro evaluation. "
                "Ours uses MaskKnownPosNeg=True and MaskSameItemNeg=True. "
                "Ours + SAGE uses tail ratio 0.01 with course feedback, course sampling, and prereq auxiliary loss applied to all items; "
                "Imp. reports the relative difference over Ours."
            ),
        },
    ]
    for spec in exports:
        require_config(spec["sage_dir"], course_scope_false=spec["course_scope_false"])
        cold_rows = build_rows(COLD_METRICS, spec["sage_dir"])
        hot_rows = build_rows(HOT_METRICS, spec["sage_dir"])
        write_csv(cold_rows, COLD_METRICS, spec["cold_csv"])
        write_csv(hot_rows, HOT_METRICS, spec["hot_csv"])
        saved_docx = build_docx(cold_rows, hot_rows, out_docx=spec["docx"], note=spec["note"])
        print(f"Wrote {saved_docx}")
        print(f"Wrote {spec['cold_csv']}")
        print(f"Wrote {spec['hot_csv']}")


if __name__ == "__main__":
    main()
