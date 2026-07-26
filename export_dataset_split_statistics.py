"""Export dataset and per-seed static split statistics.

Outputs are written to:
    output/doc/dataset_split_statistics/

The script reads the finalized balanced static item-cold split artifacts used
by the current 3-seed item-macro main table:
    outputs/content_delta_pop5/static_item_cold_balanced_itemmacro_v1/
    strict_item_cold_balanced_thr1_seed_*
"""

from __future__ import annotations

import json
import pickle
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DATA_DIR = Path("processed_data_hin_clean_pop5")
SPLIT_ROOT = Path("outputs/content_delta_pop5/static_item_cold_balanced_itemmacro_v1")
OUT_DIR = Path("output/doc/dataset_split_statistics")
SEEDS = (2025, 2026, 2027)


def load_pickle_df(path: Path) -> pd.DataFrame:
    with path.open("rb") as f:
        return pickle.load(f)


def fmt_int(value) -> str:
    return f"{int(value):,}"


def fmt_float(value: float, digits: int = 4) -> str:
    return f"{float(value):.{digits}f}"


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, spec in kwargs.items():
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in spec.items():
            element.set(qn(f"w:{key}"), str(value))


def clear_cell_borders(cell) -> None:
    nil = {"val": "nil", "sz": "0", "space": "0", "color": "FFFFFF"}
    set_cell_border(cell, top=nil, bottom=nil, left=nil, right=nil, insideH=nil, insideV=nil)


def add_rule(cell, edge: str, width: str = "6") -> None:
    set_cell_border(cell, **{edge: {"val": "single", "sz": width, "space": "0", "color": "000000"}})


def set_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=28, start=24, bottom=28, end=24) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_text(cell, text: str, *, bold=False, italic=False, size=8.4, align=None) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def style_cell(cell) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    clear_cell_borders(cell)


def apply_rules(table) -> None:
    n_rows = len(table.rows)
    n_cols = len(table.columns)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            clear_cell_borders(cell)
            if col_idx < n_cols - 1:
                add_rule(cell, "right", "5")
            if row_idx == 0:
                add_rule(cell, "top", "9")
                add_rule(cell, "bottom", "7")
            if row_idx == n_rows - 1:
                add_rule(cell, "bottom", "9")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(9.2)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(7.3)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_doc_table(doc: Document, caption: str, df: pd.DataFrame, widths: list[float], note: str = "") -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_width(cell, widths[idx])
            style_cell(cell)
    for idx, col in enumerate(df.columns):
        add_text(table.rows[0].cells[idx], col, bold=True, italic=(idx > 0), size=8.2)
    for r_idx, (_, row) in enumerate(df.iterrows(), start=1):
        for c_idx, col in enumerate(df.columns):
            add_text(table.rows[r_idx].cells[c_idx], row[col], bold=(c_idx == 0), size=8.2)
    apply_rules(table)
    if note:
        add_note(doc, note)


def load_representative_manifest() -> dict:
    path = SPLIT_ROOT / "strict_item_cold_balanced_thr1_seed_2025" / "static_protocol_manifest.json"
    if not path.exists():
        path = next(SPLIT_ROOT.glob("*/static_protocol_manifest.json"))
    return json.loads(path.read_text(encoding="utf-8"))


def build_dataset_stats() -> pd.DataFrame:
    meta = json.loads((DATA_DIR / "meta.json").read_text(encoding="utf-8"))
    df = load_pickle_df(DATA_DIR / "stream_data.pkl")
    manifest = load_representative_manifest()
    course = manifest.get("course_stats", {})
    split = manifest.get("split", {})

    rows = [
        ("Interactions", len(df)),
        ("Users", int(df["u_idx"].nunique())),
        ("Items", int(df["i_idx"].nunique())),
        ("Content Dim.", int(meta.get("content_dim", 0))),
        ("Density", len(df) / (max(1, int(df["u_idx"].nunique())) * max(1, int(df["i_idx"].nunique())))),
        ("Avg. Inter./User", len(df) / max(1, int(df["u_idx"].nunique()))),
        ("Avg. Inter./Item", len(df) / max(1, int(df["i_idx"].nunique()))),
        ("Items w/ Concepts", int(course.get("items_with_concept", 0))),
        ("Items w/ Prereq.", int(course.get("items_with_prereq", 0))),
        ("Prereq. Edges", int(course.get("prereq_edges_kept", 0))),
        ("Redundant Groups", int(course.get("redundant_family_groups", 0))),
        ("Eligible Cold Items", int(split.get("strict_item_cold_eligible_items", 0))),
    ]
    formatted = []
    for metric, value in rows:
        if isinstance(value, float):
            text = fmt_float(value, 4)
        else:
            text = fmt_int(value)
        formatted.append({"Statistic": metric, "Value": text})
    return pd.DataFrame(formatted)


def split_root(seed: int) -> Path:
    return SPLIT_ROOT / f"strict_item_cold_balanced_thr1_seed_{seed}"


def build_split_stats() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    row_rows = []
    cold_rows = []
    detail_rows = []
    for seed in SEEDS:
        root = split_root(seed)
        if not root.exists():
            raise FileNotFoundError(f"Missing split root: {root}")
        summary = json.loads((root / "static_split_summary.json").read_text(encoding="utf-8"))
        train = load_pickle_df(root / "static_train.pkl")
        val = load_pickle_df(root / "static_val.pkl")
        test = load_pickle_df(root / "static_test.pkl")

        def split_info(df: pd.DataFrame) -> dict:
            cold_mask = df["popularity"] == 0
            return {
                "rows": len(df),
                "users": int(df["u_idx"].nunique()),
                "items": int(df["i_idx"].nunique()),
                "cold_rows": int(cold_mask.sum()),
                "cold_items": int(df.loc[cold_mask, "i_idx"].nunique()),
                "hot_rows": int((~cold_mask).sum()),
            }

        ti, vi, tei = split_info(train), split_info(val), split_info(test)
        row_rows.append(
            {
                "Seed": seed,
                "Train": fmt_int(ti["rows"]),
                "Val": fmt_int(vi["rows"]),
                "Test": fmt_int(tei["rows"]),
                "Train %": fmt_float(summary["actual_train_ratio"] * 100, 2),
                "Val %": fmt_float(summary["actual_val_ratio"] * 100, 2),
                "Test %": fmt_float(summary["actual_test_ratio"] * 100, 2),
                "Val SeenU": fmt_float(summary["val_user_seen_ratio"] * 100, 1),
                "Test SeenU": fmt_float(summary["test_user_seen_ratio"] * 100, 1),
            }
        )
        cold_rows.append(
            {
                "Seed": seed,
                "Train Items": fmt_int(ti["items"]),
                "Val Items": fmt_int(vi["items"]),
                "Test Items": fmt_int(tei["items"]),
                "Val Cold Items": fmt_int(vi["cold_items"]),
                "Test Cold Items": fmt_int(tei["cold_items"]),
                "Val Cold Int.": fmt_int(vi["cold_rows"]),
                "Test Cold Int.": fmt_int(tei["cold_rows"]),
            }
        )
        detail_rows.append(
            {
                "seed": seed,
                "train_rows": ti["rows"],
                "val_rows": vi["rows"],
                "test_rows": tei["rows"],
                "actual_train_ratio": summary["actual_train_ratio"],
                "actual_val_ratio": summary["actual_val_ratio"],
                "actual_test_ratio": summary["actual_test_ratio"],
                "train_users": ti["users"],
                "val_users": vi["users"],
                "test_users": tei["users"],
                "train_items": ti["items"],
                "val_items": vi["items"],
                "test_items": tei["items"],
                "train_cold_rows": ti["cold_rows"],
                "val_cold_rows": vi["cold_rows"],
                "test_cold_rows": tei["cold_rows"],
                "val_hot_rows": vi["hot_rows"],
                "test_hot_rows": tei["hot_rows"],
                "val_cold_items": vi["cold_items"],
                "test_cold_items": tei["cold_items"],
                "val_cold_item_pop_sum": summary["strict_item_cold_val_item_pop_sum"],
                "test_cold_item_pop_sum": summary["strict_item_cold_test_item_pop_sum"],
                "val_cold_item_pop_mean": summary["strict_item_cold_val_item_pop_mean"],
                "test_cold_item_pop_mean": summary["strict_item_cold_test_item_pop_mean"],
                "val_fold_ids": ",".join(str(x) for x in summary["strict_item_cold_val_fold_ids"]),
                "test_fold_ids": ",".join(str(x) for x in summary["strict_item_cold_test_fold_ids"]),
                "train_item_coverage_moves": summary["train_item_coverage_moves"],
            }
        )
    return pd.DataFrame(row_rows), pd.DataFrame(cold_rows), pd.DataFrame(detail_rows)


def write_latex(dataset_df: pd.DataFrame, split_df: pd.DataFrame, cold_df: pd.DataFrame) -> str:
    def tabular(df: pd.DataFrame, caption: str, label: str) -> str:
        cols = "l" + "c" * (len(df.columns) - 1)
        lines = [
            "\\begin{table}[t]",
            "\\centering",
            "\\small",
            f"\\caption{{{caption}}}",
            f"\\label{{{label}}}",
            f"\\begin{{tabular}}{{{cols}}}",
            "\\toprule",
            " & ".join(df.columns) + " \\\\",
            "\\midrule",
        ]
        for _, row in df.iterrows():
            lines.append(" & ".join(str(row[col]) for col in df.columns) + " \\\\")
        lines.extend(["\\bottomrule", "\\end{tabular}", "\\end{table}", ""])
        return "\n".join(lines)

    return "\n".join(
        [
            tabular(dataset_df, "Dataset statistics.", "tab:dataset-statistics"),
            tabular(split_df, "Per-seed interaction split statistics.", "tab:split-statistics"),
            tabular(cold_df, "Per-seed item-cold split statistics.", "tab:cold-split-statistics"),
        ]
    )


def build_docx(dataset_df: pd.DataFrame, split_df: pd.DataFrame, cold_df: pd.DataFrame, out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(9)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_doc_table(
        doc,
        "Table A: Dataset statistics.",
        dataset_df,
        widths=[1.55, 1.1],
        note="Statistics are computed from processed_data_hin_clean_pop5. Density is interactions divided by users times items.",
    )
    add_doc_table(
        doc,
        "Table B: Per-seed interaction split statistics.",
        split_df,
        widths=[0.62, 0.76, 0.76, 0.76, 0.62, 0.62, 0.62, 0.72, 0.72],
        note="SeenU denotes the percentage of validation/test users that appear in the training set.",
    )
    add_doc_table(
        doc,
        "Table C: Per-seed item-cold split statistics.",
        cold_df,
        widths=[0.62, 0.74, 0.70, 0.70, 0.82, 0.86, 0.82, 0.86],
        note="Cold items are defined as items with zero training interactions; Train has no cold rows by construction.",
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    dataset_df = build_dataset_stats()
    split_df, cold_df, detail_df = build_split_stats()

    dataset_df.to_csv(OUT_DIR / "dataset_statistics.csv", index=False)
    split_df.to_csv(OUT_DIR / "split_statistics_by_seed.csv", index=False)
    cold_df.to_csv(OUT_DIR / "cold_split_statistics_by_seed.csv", index=False)
    detail_df.to_csv(OUT_DIR / "split_statistics_by_seed_detail.csv", index=False)
    (OUT_DIR / "dataset_split_statistics.tex").write_text(
        write_latex(dataset_df, split_df, cold_df),
        encoding="utf-8",
    )
    build_docx(dataset_df, split_df, cold_df, OUT_DIR / "dataset_split_statistics.docx")
    print(f"Wrote {OUT_DIR / 'dataset_statistics.csv'}")
    print(f"Wrote {OUT_DIR / 'split_statistics_by_seed.csv'}")
    print(f"Wrote {OUT_DIR / 'cold_split_statistics_by_seed.csv'}")
    print(f"Wrote {OUT_DIR / 'split_statistics_by_seed_detail.csv'}")
    print(f"Wrote {OUT_DIR / 'dataset_split_statistics.tex'}")
    print(f"Wrote {OUT_DIR / 'dataset_split_statistics.docx'}")


if __name__ == "__main__":
    main()
