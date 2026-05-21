"""Export final narrow top-conference tables for the static item-cold study.

Outputs:
  - DOCX with compact main and ablation tables.
  - CSV files used by the paper tables.
  - LaTeX booktabs table fragments.

The main table uses the corrected E60 course-aware run:
outputs/content_delta_pop5/course_ablation_e60_3seed_corrected/full
Override with USIM_COURSE_ABLATION_ROOT when exporting legacy runs.
"""

from __future__ import annotations

from datetime import datetime
import os
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT_DIR = Path("output/doc/final_narrow_topconf")
BASELINE_CSV = Path(
    "outputs/content_delta_pop5/static_item_cold_balanced/"
    "main_table_item_macro_final_audit_with_dropoutnet_official_teacher80_student120_cgrc_paper/"
    "main_table_item_macro_summary.csv"
)
ABLATION_ROOT = Path(
    os.environ.get(
        "USIM_COURSE_ABLATION_ROOT",
        "outputs/content_delta_pop5/course_ablation_e60_3seed_corrected",
    )
)
FAST3_CSV = ABLATION_ROOT / "full" / "fast3_static_multiseed_summary.csv"

METRICS = [
    ("Cold R@5", "cold_r5"),
    ("Cold R@10", "cold_r10"),
    ("Cold R@20", "cold_r20"),
    ("Cold N@5", "cold_n5"),
    ("Cold N@10", "cold_n10"),
    ("Cold N@20", "cold_n20"),
]

MODEL_ORDER = [
    "Popularity",
    "BPR",
    "LightGCN",
    "DropoutNet",
    "ContentProfile",
    "CCFCRec",
    "ALDI",
    "CGRC",
    "Ours",
]

MODEL_RENAME = {
    "CGRC-paper": "CGRC",
}

ABLATION_ORDER = [
    ("Ours", ABLATION_ROOT / "full" / "fast3_static_multiseed_summary.csv"),
    ("w/o Course-aware Reward", ABLATION_ROOT / "wo_course_reward" / "fast3_static_multiseed_summary.csv"),
    ("w/o Course-aware User Selection", ABLATION_ROOT / "wo_course_candidate" / "fast3_static_multiseed_summary.csv"),
    ("w/o Prereq. Auxiliary Loss", ABLATION_ROOT / "wo_prereq_aux" / "fast3_static_multiseed_summary.csv"),
    ("w/o All Course Signals", ABLATION_ROOT / "wo_all_course_signals" / "fast3_static_multiseed_summary.csv"),
]


def metric_value(row: pd.Series, metric_id: str) -> float:
    split, metric = metric_id.split("_", 1)
    return float(row[f"full_{split}_item_macro_{metric}_mean"])


def metric_std(row: pd.Series, metric_id: str) -> float:
    split, metric = metric_id.split("_", 1)
    return float(row[f"full_{split}_item_macro_{metric}_std"])


def baseline_metric(row: pd.Series, metric_id: str, suffix: str) -> float:
    split, metric = metric_id.split("_", 1)
    return float(row[f"{split}_{metric.upper()}_{suffix}"])


def load_main_table() -> pd.DataFrame:
    base = pd.read_csv(BASELINE_CSV)
    rows = []
    for _, row in base.iterrows():
        model_name = MODEL_RENAME.get(str(row["model"]), str(row["model"]))
        out = {
            "Model": model_name,
            "runs": row.get("runs", ""),
            "seeds": row.get("seeds", ""),
        }
        for label, metric_id in METRICS:
            out[label] = baseline_metric(row, metric_id, "mean")
            out[f"{label} std"] = baseline_metric(row, metric_id, "std")
        rows.append(out)

    fast3 = pd.read_csv(FAST3_CSV).iloc[0]
    rows.append(
        {
            "Model": "Ours",
            "runs": fast3.get("runs", ""),
            "seeds": fast3.get("seeds", ""),
        }
    )
    for label, metric_id in METRICS:
        rows[-1][label] = metric_value(fast3, metric_id)
        rows[-1][f"{label} std"] = metric_std(fast3, metric_id)

    df = pd.DataFrame(rows)
    order = {name: idx for idx, name in enumerate(MODEL_ORDER)}
    df["__order"] = df["Model"].map(order).fillna(999)
    return df.sort_values(["__order", "Model"]).drop(columns=["__order"]).reset_index(drop=True)


def load_ablation_table() -> pd.DataFrame:
    rows = []
    for name, path in ABLATION_ORDER:
        if not path.exists():
            continue
        row = pd.read_csv(path).iloc[0]
        rows.append(
            {
                "Model": name,
                "runs": row.get("runs", ""),
                "seeds": row.get("seeds", ""),
            }
        )
        for label, metric_id in METRICS:
            rows[-1][label] = metric_value(row, metric_id)
            rows[-1][f"{label} std"] = metric_std(row, metric_id)
    return pd.DataFrame(rows)


def best_second(df: pd.DataFrame) -> dict[str, tuple[float | None, float | None]]:
    ranks = {}
    for label, _ in METRICS:
        values = sorted({float(v) for v in df[label].dropna().tolist()}, reverse=True)
        ranks[label] = (values[0] if values else None, values[1] if len(values) > 1 else None)
    return ranks


def fmt(value: float) -> str:
    return f"{float(value):.4f}"


def fmt_imp(value: float | None) -> str:
    if value is None:
        return "NA"
    sign = "+" if value >= 0 else ""
    return f"{sign}{value * 100:.2f}%"


def improvement_row(df: pd.DataFrame, ours: str = "Ours") -> dict[str, str]:
    ours_row = df[df["Model"] == ours]
    if ours_row.empty:
        return {label: "NA" for label, _ in METRICS}
    ours_row = ours_row.iloc[0]
    base = df[df["Model"] != ours]
    out = {}
    for label, _ in METRICS:
        best_base = float(base[label].max())
        out[label] = fmt_imp(None if abs(best_base) < 1e-12 else (float(ours_row[label]) - best_base) / best_base)
    return out


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, spec in kwargs.items():
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in spec.items():
            element.set(qn(f"w:{key}"), str(value))


def clear_cell_borders(cell) -> None:
    nil = {"val": "nil", "sz": "0", "space": "0", "color": "FFFFFF"}
    set_cell_border(cell, top=nil, bottom=nil, left=nil, right=nil, insideH=nil, insideV=nil)


def add_rule(cell, edge: str, width: str = "6") -> None:
    set_cell_border(cell, **{edge: {"val": "single", "sz": width, "space": "0", "color": "000000"}})


def set_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=32, start=24, bottom=32, end=24) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_text(cell, text: str, *, bold=False, underline=False, italic=False, size=8.7, align=None) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.underline = underline
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def style_cell(cell) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    clear_cell_borders(cell)


def apply_rules(table, separator_rows: tuple[int, ...] = ()) -> None:
    n_rows = len(table.rows)
    n_cols = len(table.columns)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            clear_cell_borders(cell)
            if col_idx < n_cols - 1:
                add_rule(cell, "right", "5")
            if row_idx == 0:
                add_rule(cell, "top", "9")
                add_rule(cell, "bottom", "7")
            if row_idx in separator_rows:
                add_rule(cell, "bottom", "6")
            if row_idx == n_rows - 1:
                add_rule(cell, "bottom", "9")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(9.2)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(7.5)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_table(doc: Document, df: pd.DataFrame, *, include_imp: bool, caption: str, note: str) -> None:
    add_caption(doc, caption)
    ranks = best_second(df)
    n_extra = 1 if include_imp else 0
    table = doc.add_table(rows=len(df) + 1 + n_extra, cols=1 + len(METRICS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.20] + [0.58] * len(METRICS)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_width(cell, widths[idx])
            style_cell(cell)

    headers = ["Model"] + [label for label, _ in METRICS]
    for idx, header in enumerate(headers):
        add_text(table.rows[0].cells[idx], header, bold=(idx == 0), italic=(idx > 0), size=7.8)

    for row_idx, (_, row) in enumerate(df.iterrows(), start=1):
        model = str(row["Model"])
        add_text(table.rows[row_idx].cells[0], model, bold=True, italic=(model == "Ours"), size=8.2)
        for col_idx, (label, _) in enumerate(METRICS, start=1):
            value = float(row[label])
            best, second = ranks[label]
            is_best = best is not None and abs(value - best) < 5e-13
            is_second = second is not None and abs(value - second) < 5e-13
            add_text(table.rows[row_idx].cells[col_idx], fmt(value), bold=is_best, underline=is_second, size=8.0)

    if include_imp:
        imp = improvement_row(df)
        row_idx = len(df) + 1
        add_text(table.rows[row_idx].cells[0], "Imp.", bold=True, size=8.7)
        for col_idx, (label, _) in enumerate(METRICS, start=1):
            add_text(table.rows[row_idx].cells[col_idx], imp[label], bold=True, size=7.8)

    separator_rows = []
    if "LightGCN" in set(df["Model"]):
        separator_rows.append(int(df.index[df["Model"] == "LightGCN"][0]) + 1)
    if "Ours" in set(df["Model"]):
        separator_rows.append(int(df.index[df["Model"] == "Ours"][0]))
    apply_rules(table, tuple(separator_rows))
    add_note(doc, note)


def csv_for_paper(df: pd.DataFrame, path: Path, *, include_imp: bool = False) -> None:
    rows = []
    for _, row in df.iterrows():
        out = {"Model": row["Model"]}
        for label, _ in METRICS:
            out[label] = fmt(row[label])
        rows.append(out)
    if include_imp:
        imp = improvement_row(df)
        out = {"Model": "Imp."}
        for label, _ in METRICS:
            out[label] = imp[label]
        rows.append(out)
    pd.DataFrame(rows).to_csv(path, index=False)


def latex_table(df: pd.DataFrame, *, include_imp: bool, caption: str, label: str) -> str:
    ranks = best_second(df)
    lines = [
        "\\begin{table}[t]",
        "\\centering",
        "\\small",
        f"\\caption{{{caption}}}",
        f"\\label{{{label}}}",
        "\\begin{tabular}{l" + "c" * len(METRICS) + "}",
        "\\toprule",
        "Model & " + " & ".join(label for label, _ in METRICS) + " \\\\",
        "\\midrule",
    ]
    for _, row in df.iterrows():
        vals = []
        for metric, _ in METRICS:
            value = float(row[metric])
            best, second = ranks[metric]
            text = fmt(value)
            if best is not None and abs(value - best) < 5e-13:
                text = f"\\textbf{{{text}}}"
            elif second is not None and abs(value - second) < 5e-13:
                text = f"\\underline{{{text}}}"
            vals.append(text)
        lines.append(f"{row['Model']} & " + " & ".join(vals) + " \\\\")
    if include_imp:
        imp = improvement_row(df)
        lines.extend([
            "\\midrule",
            "Imp. & " + " & ".join(imp[m].replace("%", "\\%") for m, _ in METRICS) + " \\\\",
        ])
    lines.extend(["\\bottomrule", "\\end{tabular}", "\\end{table}", ""])
    return "\n".join(lines)


def build_docx(main_df: pd.DataFrame, ablation_df: pd.DataFrame, out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(9)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_table(
        doc,
        main_df,
        include_imp=True,
        caption=(
            "Table 1: Cold-item performance comparison on the balanced static item-cold split. "
            "The best result is in bold and the second-best result is underlined."
        ),
        note=(
            "All values are three-seed means under full-ranking item-macro evaluation."
        ),
    )
    add_table(
        doc,
        ablation_df,
        include_imp=False,
        caption="Table 2: Cold-item ablation study of Ours under the same protocol.",
        note="All values are three-seed means. The main method is Ours without pseudo-cold training.",
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    save_docx_with_fallback(doc, out_path)


def build_single_table_docx(
    df: pd.DataFrame,
    out_path: Path,
    *,
    caption: str,
    note: str,
    include_imp: bool = False,
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(9)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_table(
        doc,
        df,
        include_imp=include_imp,
        caption=caption,
        note=note,
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    save_docx_with_fallback(doc, out_path)


def save_docx_with_fallback(doc: Document, out_path: Path) -> Path:
    try:
        doc.save(out_path)
        return out_path
    except PermissionError:
        fallback = out_path.with_name(f"{out_path.stem}_cold_all{out_path.suffix}")
        try:
            doc.save(fallback)
            print(f"Locked: {out_path}; wrote fallback {fallback}")
            return fallback
        except PermissionError:
            stamped = out_path.with_name(
                f"{out_path.stem}_cold_all_{datetime.now().strftime('%Y%m%d_%H%M%S')}{out_path.suffix}"
            )
            doc.save(stamped)
            print(f"Locked: {out_path} and {fallback}; wrote fallback {stamped}")
            return stamped


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    main_df = load_main_table()
    ablation_df = load_ablation_table()

    csv_for_paper(main_df, OUT_DIR / "main_table_narrow_topconf.csv", include_imp=True)
    csv_for_paper(ablation_df, OUT_DIR / "ablation_table_narrow_topconf.csv")

    tex = latex_table(
        main_df,
        include_imp=True,
        caption=(
            "Cold-item performance comparison on the balanced static item-cold split. "
            "The best result is in bold and the second-best result is underlined."
        ),
        label="tab:main-item-cold",
    )
    tex += "\n" + latex_table(
        ablation_df,
        include_imp=False,
        caption="Cold-item ablation study of Ours under the same protocol.",
        label="tab:ablation-fast3",
    )
    (OUT_DIR / "tables_narrow_topconf.tex").write_text(tex, encoding="utf-8")
    build_docx(main_df, ablation_df, OUT_DIR / "tables_narrow_topconf.docx")
    build_single_table_docx(
        main_df,
        OUT_DIR / "main_table_narrow_topconf_cold_all.docx",
        caption=(
            "Table 1: Cold-item performance comparison on the balanced static item-cold split. "
            "The best result is in bold and the second-best result is underlined."
        ),
        note="All values are three-seed means under full-ranking item-macro evaluation.",
        include_imp=True,
    )
    build_single_table_docx(
        ablation_df,
        OUT_DIR / "ablation_table_narrow_topconf.docx",
        caption="Table 2: Cold-item ablation study of Ours under the same protocol.",
        note="All values are three-seed means under full-ranking item-macro evaluation.",
    )

    print(f"Wrote {OUT_DIR / 'main_table_narrow_topconf.csv'}")
    print(f"Wrote {OUT_DIR / 'ablation_table_narrow_topconf.csv'}")
    print(f"Wrote {OUT_DIR / 'tables_narrow_topconf.tex'}")
    print(f"Wrote {OUT_DIR / 'tables_narrow_topconf.docx'}")
    print(f"Wrote {OUT_DIR / 'main_table_narrow_topconf_cold_all.docx'}")
    print(f"Wrote {OUT_DIR / 'ablation_table_narrow_topconf.docx'}")


if __name__ == "__main__":
    main()
