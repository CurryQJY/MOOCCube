"""Export the 3-seed @5/@10/@20 main table as split Cold/Hot Word tables.

Default output follows a compact top-conference table style and avoids merged
metric headers so Word border lines remain continuous.

Run:
    .\\py.bat export_main_table_3seed_docx.py
"""

from __future__ import annotations

import argparse
from pathlib import Path
from shutil import copyfile

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DEFAULT_BASELINE_CSV = Path(
    "outputs/content_delta_pop5/static_item_cold_balanced/"
    "main_table_item_macro_final_audit/main_table_item_macro_summary.csv"
)
DEFAULT_FAST3_CSV = Path(
    "outputs/content_delta_pop5/static_item_cold_balanced_itemmacro_v1/"
    "fast3_static_multiseed_summary.csv"
)
DEFAULT_OUT = Path("output/doc/main_table_3seed_item_macro_split_full_topconf.docx")
DEFAULT_COMPAT_OUT = Path("output/doc/main_table_3seed_item_macro_reference_topconf.docx")
DEFAULT_MEAN_STD_OUT = Path("output/doc/main_table_3seed_item_macro_split_full_topconf_mean_std.docx")

MODEL_ORDER = [
    "Popularity",
    "BPR",
    "LightGCN",
    "DropoutNet",
    "ContentProfile",
    "CCFCRec",
    "ALDI",
    "CGRC-paper",
    "FAST3",
]

SPLITS = ("Cold", "Hot")
KS = (5, 10, 20)
METRICS = ("Recall", "NDCG")
TABLE_COLS = ["Model"] + [f"{metric}@{k}" for metric in METRICS for k in KS]


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge, spec in kwargs.items():
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in spec.items():
            element.set(qn("w:{}".format(key)), str(value))


def clear_cell_borders(cell) -> None:
    nil = {"val": "nil", "sz": "0", "space": "0", "color": "FFFFFF"}
    set_cell_border(cell, top=nil, bottom=nil, left=nil, right=nil, insideH=nil, insideV=nil)


def add_rule(cell, edge: str, width: str = "6", color: str = "000000") -> None:
    set_cell_border(cell, **{edge: {"val": "single", "sz": width, "space": "0", "color": color}})


def set_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=24, start=20, bottom=24, end=20) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:{}".format(margin)))
        if node is None:
            node = OxmlElement("w:{}".format(margin))
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_height(row, height_twips: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(height_twips))
    tr_height.set(qn("w:hRule"), "atLeast")


def add_text(
    cell,
    text: str,
    *,
    bold=False,
    underline=False,
    italic=False,
    size=8.6,
    align=None,
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.underline = underline
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def style_cell(cell) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    clear_cell_borders(cell)


def metric_key(split: str, metric: str, k: int) -> str:
    prefix = split.lower()
    tag = "R" if metric == "Recall" else "N"
    return f"{prefix}_{tag}{k}"


def format_mean(value) -> str:
    return "{:.4f}".format(float(value))


def format_mean_std(mean, std) -> str:
    return "{:.4f}\n(+/-{:.4f})".format(float(mean), float(std))


def format_imp(value) -> str:
    sign = "+" if value >= 0 else ""
    return "{}{:.2f}%".format(sign, value * 100.0)


def seed_text(df: pd.DataFrame) -> str:
    if "seeds" not in df.columns:
        return "2025, 2026, and 2027"
    values = [str(x) for x in df["seeds"].dropna().unique().tolist()]
    if not values:
        return "2025, 2026, and 2027"
    parts = values[0].replace(",", ", ").split(", ")
    return parts[0] if len(parts) == 1 else ", ".join(parts[:-1]) + ", and " + parts[-1]


def load_baseline_rows(path: Path) -> pd.DataFrame:
    raw = pd.read_csv(path)
    rows = []
    for _, row in raw.iterrows():
        out = {
            "Model": str(row["model"]),
            "runs": row.get("runs"),
            "seeds": row.get("seeds"),
        }
        for split in SPLITS:
            src_prefix = split.lower()
            for metric in METRICS:
                tag = "R" if metric == "Recall" else "N"
                for k in KS:
                    key = metric_key(split, metric, k)
                    out[f"{key}_mean"] = row[f"{src_prefix}_{tag}{k}_mean"]
                    out[f"{key}_std"] = row[f"{src_prefix}_{tag}{k}_std"]
        rows.append(out)
    return pd.DataFrame(rows)


def load_fast3_row(path: Path) -> pd.DataFrame:
    raw = pd.read_csv(path)
    if raw.empty:
        raise ValueError(f"FAST3 summary is empty: {path}")
    row = raw.iloc[0]
    out = {
        "Model": "FAST3",
        "runs": row.get("runs"),
        "seeds": row.get("seeds"),
    }
    for split in SPLITS:
        src_prefix = f"full_{split.lower()}_item_macro"
        for metric in METRICS:
            tag = "r" if metric == "Recall" else "n"
            for k in KS:
                key = metric_key(split, metric, k)
                out[f"{key}_mean"] = row[f"{src_prefix}_{tag}{k}_mean"]
                out[f"{key}_std"] = row[f"{src_prefix}_{tag}{k}_std"]
    return pd.DataFrame([out])


def load_combined_table(baseline_csv: Path, fast3_csv: Path) -> pd.DataFrame:
    if not baseline_csv.exists():
        raise FileNotFoundError(f"Missing baseline summary CSV: {baseline_csv}")
    if not fast3_csv.exists():
        raise FileNotFoundError(f"Missing FAST3 summary CSV: {fast3_csv}")
    df = pd.concat([load_baseline_rows(baseline_csv), load_fast3_row(fast3_csv)], ignore_index=True)
    order = {name: idx for idx, name in enumerate(MODEL_ORDER)}
    df["__order"] = df["Model"].map(order).fillna(999)
    return df.sort_values(["__order", "Model"]).drop(columns=["__order"])


def best_and_second(df: pd.DataFrame, split: str) -> dict:
    ranks = {}
    for metric in METRICS:
        for k in KS:
            key = metric_key(split, metric, k)
            vals = sorted({float(v) for v in df[f"{key}_mean"].dropna().tolist()}, reverse=True)
            ranks[key] = (vals[0] if vals else None, vals[1] if len(vals) > 1 else None)
    return ranks


def improvement_values(df: pd.DataFrame, split: str, ours_name: str = "FAST3") -> dict:
    ours = df[df["Model"] == ours_name]
    if ours.empty:
        return {}
    ours = ours.iloc[0]
    baselines = df[df["Model"] != ours_name]
    out = {}
    for metric in METRICS:
        for k in KS:
            key = metric_key(split, metric, k)
            base = float(baselines[f"{key}_mean"].max())
            out[key] = None if abs(base) < 1e-12 else (float(ours[f"{key}_mean"]) - base) / base
    return out


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.2)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_split_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.space_before = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(9.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def apply_continuous_borders(table, separator_rows: tuple[int, ...]) -> None:
    for row in table.rows:
        for cell in row.cells:
            clear_cell_borders(cell)

    n_rows = len(table.rows)
    n_cols = len(table.columns)

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if col_idx < n_cols - 1:
                add_rule(cell, "right", width="5")
            if row_idx == 0:
                add_rule(cell, "top", width="10")
                add_rule(cell, "bottom", width="8")
            if row_idx in separator_rows:
                add_rule(cell, "bottom", width="7")
            if row_idx == n_rows - 1:
                add_rule(cell, "bottom", width="10")


def build_split_table(doc: Document, df: pd.DataFrame, split: str, *, include_std: bool = False) -> None:
    add_split_title(doc, f"{split} results")
    ranks = best_and_second(df, split)
    imp = improvement_values(df, split)

    table = doc.add_table(rows=len(df) + 2, cols=len(TABLE_COLS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.18] + ([0.76] * 6 if include_std else [0.68] * 6)

    for row in table.rows:
        set_row_height(row, 390 if include_std else 285)
        for idx, width in enumerate(widths):
            set_width(row.cells[idx], width)
            style_cell(row.cells[idx])

    header = table.rows[0].cells
    for idx, label in enumerate(TABLE_COLS):
        add_text(
            header[idx],
            label,
            bold=(idx == 0),
            italic=(idx > 0),
            size=8.8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    set_repeat_table_header(table.rows[0])

    out_idx = 1
    for _, row in df.iterrows():
        cells = table.rows[out_idx].cells
        model = str(row["Model"])
        add_text(cells[0], model, bold=True, italic=(model == "FAST3"), size=8.8)
        col_idx = 1
        for metric in METRICS:
            for k in KS:
                key = metric_key(split, metric, k)
                mean = float(row[f"{key}_mean"])
                std = float(row[f"{key}_std"])
                best, second = ranks[key]
                is_best = best is not None and abs(mean - best) < 5e-13
                is_second = second is not None and abs(mean - second) < 5e-13
                text = format_mean_std(mean, std) if include_std else format_mean(mean)
                add_text(
                    cells[col_idx],
                    text,
                    bold=is_best,
                    underline=is_second,
                    size=7.1 if include_std else 8.8,
                )
                col_idx += 1
        out_idx += 1

    cells = table.rows[out_idx].cells
    add_text(cells[0], "Imp.", bold=True, size=8.8)
    col_idx = 1
    for metric in METRICS:
        for k in KS:
            key = metric_key(split, metric, k)
            val = imp.get(key)
            add_text(cells[col_idx], "NA" if val is None else format_imp(val), bold=True, size=8.6)
            col_idx += 1

    # Header row = 0. Separator after LightGCN and before FAST3/Imp.
    separator_rows = (3, len(df) - 1)
    apply_continuous_borders(table, separator_rows)


def build_docx(df: pd.DataFrame, out_path: Path, *, include_std: bool = False) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.42)
    section.right_margin = Inches(0.42)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(9)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    caption = (
        "Table 1: Performance comparison on the balanced static item-cold split with K = 5, 10, and 20. "
        "The best performance is indicated in bold, while the second-best performance is underlined "
        f"(three-seed average over {seed_text(df)}; the relative improvements over the best baseline "
        "are denoted as Imp.)."
    )
    if include_std:
        caption += " Values are reported as mean +/- std."
    add_caption(doc, caption)

    build_split_table(doc, df, "Cold", include_std=include_std)
    build_split_table(doc, df, "Hot", include_std=include_std)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def try_build_docx(df: pd.DataFrame, out_path: Path, *, include_std: bool = False) -> bool:
    try:
        build_docx(df, out_path, include_std=include_std)
    except PermissionError:
        print(f"Skip overwrite locked {out_path}")
        return False
    print(f"Wrote {out_path}")
    return True


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--baseline-csv", type=Path, default=DEFAULT_BASELINE_CSV)
    parser.add_argument("--fast3-csv", type=Path, default=DEFAULT_FAST3_CSV)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    parser.add_argument("--compat-out", type=Path, default=DEFAULT_COMPAT_OUT)
    parser.add_argument("--mean-std-out", type=Path, default=DEFAULT_MEAN_STD_OUT)
    parser.add_argument("--no-mean-std", action="store_true")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    df = load_combined_table(args.baseline_csv, args.fast3_csv)
    wrote_main = try_build_docx(df, args.out, include_std=False)
    if wrote_main:
        try:
            args.compat_out.parent.mkdir(parents=True, exist_ok=True)
            copyfile(args.out, args.compat_out)
            print(f"Wrote {args.compat_out}")
        except PermissionError:
            print(f"Skip overwrite locked {args.compat_out}")
    if not args.no_mean_std:
        try_build_docx(df, args.mean_std_out, include_std=True)


if __name__ == "__main__":
    main()
