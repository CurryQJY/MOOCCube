"""Export a top-conference style dataset statistics table to Word.

The generated table follows a compact three-line format:
caption, bold header, top/header/bottom horizontal rules only.
"""

from __future__ import annotations

import json
import pickle
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DATA_DIR = Path("processed_data_hin_clean_pop5")
OUT_DIR = Path("output/doc/dataset_topconf_table")
OUT_PATH = OUT_DIR / "dataset_statistics_topconf.docx"


def fmt_int(value: int) -> str:
    return f"{int(value):,}"


def fmt_percent(ratio: float) -> str:
    return f"{ratio * 100:.4f}%"


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, spec in kwargs.items():
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in spec.items():
            element.set(qn(f"w:{key}"), str(value))


def clear_cell_borders(cell) -> None:
    nil = {"val": "nil", "sz": "0", "space": "0", "color": "FFFFFF"}
    set_cell_border(cell, top=nil, bottom=nil, left=nil, right=nil)


def add_rule(cell, edge: str, width: str) -> None:
    set_cell_border(cell, **{edge: {"val": "single", "sz": width, "space": "0", "color": "000000"}})


def set_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=42, start=28, bottom=42, end=28) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_text(cell, text: str, *, bold=False, size=11.0) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def apply_three_line_rules(table) -> None:
    n_rows = len(table.rows)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            clear_cell_borders(cell)
            if row_idx == 0:
                add_rule(cell, "top", "10")
                add_rule(cell, "bottom", "7")
            if row_idx == n_rows - 1:
                add_rule(cell, "bottom", "10")


def load_dataset_stats() -> pd.DataFrame:
    with (DATA_DIR / "stream_data.pkl").open("rb") as f:
        df = pickle.load(f)
    meta_path = DATA_DIR / "meta.json"
    dataset_name = "MOOCCube"
    if meta_path.exists():
        meta = json.loads(meta_path.read_text(encoding="utf-8"))
        dataset_name = meta.get("dataset", dataset_name)

    users = int(df["u_idx"].nunique())
    items = int(df["i_idx"].nunique())
    interactions = int(len(df))
    density = interactions / max(1, users * items)
    return pd.DataFrame(
        [
            {
                "Datasets": dataset_name,
                "#Users": fmt_int(users),
                "#Items": fmt_int(items),
                "#Interactions": fmt_int(interactions),
                "Density": fmt_percent(density),
            }
        ]
    )


def build_docx(df: pd.DataFrame, out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(18)
    run = caption.add_run("Table 1: The statistics of the dataset.")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [1.75, 1.1, 1.1, 1.55, 1.1]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_width(cell, widths[idx])
            set_cell_margins(cell)
            clear_cell_borders(cell)

    for idx, col in enumerate(df.columns):
        add_text(table.rows[0].cells[idx], col, bold=True, size=11.0)
    for r_idx, (_, row) in enumerate(df.iterrows(), start=1):
        for c_idx, col in enumerate(df.columns):
            add_text(table.rows[r_idx].cells[c_idx], row[col], bold=(c_idx == 0), size=11.0)

    apply_three_line_rules(table)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    df = load_dataset_stats()
    build_docx(df, OUT_PATH)
    csv_path = OUT_PATH.with_suffix(".csv")
    df.to_csv(csv_path, index=False, encoding="utf-8-sig")
    print(f"Wrote {OUT_PATH}")
    print(f"Wrote {csv_path}")


if __name__ == "__main__":
    main()
