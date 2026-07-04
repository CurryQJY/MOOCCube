import csv
import json
import math
import statistics
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"D:\DeskTop\MOOCCube")
MAIN_SUMMARY = ROOT / "outputs" / "junyi" / "main_table_3seed" / "junyi_main_table_3seed_combined_summary.csv"
OUT_DIR = ROOT / "output" / "doc" / "junyi_results"
DOCX_PATH = OUT_DIR / "junyi_results_tables.docx"
CSV_PATH = OUT_DIR / "junyi_main_table_strict_tt.csv"
MASK_CSV_PATH = OUT_DIR / "junyi_mask_seed2025_table.csv"

MODEL_ORDER = [
    ("Popularity", "Popularity"),
    ("BPR", "BPR"),
    ("LightGCN", "LightGCN"),
    ("Content-CBF", "ContentProfile"),
    ("DropoutNet", "DropoutNet"),
    ("CCFCRec", "CCFCRec"),
    ("ALDI", "ALDI"),
    ("CGRC-paper", "CGRC-paper"),
    ("Ours", "Ours"),
]

MAIN_METRICS = [
    ("Cold Item R@10", "cold_item_R10"),
    ("Cold Item N@10", "cold_item_N10"),
    ("Cold Item R@20", "cold_item_R20"),
    ("Cold Item N@20", "cold_item_N20"),
    ("Hot Item N@10", "hot_item_N10"),
]

OURS_TT_RUNS = [
    ROOT / "outputs" / "junyi" / "mask_ablation" / "mask_tt" / "strict_item_cold_balanced_thr1_seed_2025",
    ROOT / "outputs" / "junyi" / "main_table_3seed" / "strict_item_cold_balanced_thr1_seed_2026",
    ROOT / "outputs" / "junyi" / "main_table_3seed" / "strict_item_cold_balanced_thr1_seed_2027",
]

MASK_RUNS = [
    ("False/True (previous)", ROOT / "outputs" / "junyi" / "official_prereq_seed2025" / "strict_item_cold_balanced_thr1_seed_2025"),
    ("True/True", ROOT / "outputs" / "junyi" / "mask_ablation" / "mask_tt" / "strict_item_cold_balanced_thr1_seed_2025"),
    ("False/False", ROOT / "outputs" / "junyi" / "mask_ablation" / "mask_ff" / "strict_item_cold_balanced_thr1_seed_2025"),
]

FINAL_FULLRANK = "final_fullrank_usim_feedback_fast3_content_delta_static.csv"
MANIFEST = "static_protocol_manifest.json"


def read_csv_row(path: Path) -> dict:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return next(csv.DictReader(f))


def read_manifest(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def mean_std(values: list[float]) -> tuple[float, float]:
    if len(values) == 1:
        return values[0], 0.0
    return sum(values) / len(values), statistics.stdev(values)


def fmt_mean_std(mean: float, std: float) -> str:
    return f"{mean:.4f} +/- {std:.4f}"


def fmt_value(value: float) -> str:
    return f"{value:.4f}"


def load_main_summary() -> dict[str, dict]:
    rows = {}
    with MAIN_SUMMARY.open("r", encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            rows[row["model"]] = row
    return rows


def assert_true_true_ours() -> dict[str, tuple[float, float]]:
    mapped = {
        "cold_item_R10": "full_cold_item_macro_r10",
        "cold_item_N10": "full_cold_item_macro_n10",
        "cold_item_R20": "full_cold_item_macro_r20",
        "cold_item_N20": "full_cold_item_macro_n20",
        "hot_item_N10": "full_hot_item_macro_n10",
    }
    detail = []
    for run_dir in OURS_TT_RUNS:
        manifest = read_manifest(run_dir / MANIFEST)
        cfg = manifest.get("model_config", {})
        if cfg.get("mask_known_pos_neg") is not True or cfg.get("mask_same_item_neg") is not True:
            raise ValueError(f"Ours True/True manifest mismatch: {run_dir}")
        if cfg.get("use_content_delta") is not False:
            raise ValueError(f"Ours expected UseContentDelta=False: {run_dir}")
        detail.append(read_csv_row(run_dir / FINAL_FULLRANK))

    out = {}
    for out_key, csv_key in mapped.items():
        values = [float(row[csv_key]) for row in detail]
        out[out_key] = mean_std(values)
    return out


def build_main_rows() -> list[dict]:
    source = load_main_summary()
    strict_ours = assert_true_true_ours()
    rows = []
    for model, source_model in MODEL_ORDER:
        if source_model not in source and model != "Ours":
            continue
        row = {
            "Model": model,
            "n": "3",
            "Seeds": "2025;2026;2027",
        }
        if model == "Ours":
            for _, key in MAIN_METRICS:
                mean, std = strict_ours[key]
                row[key] = fmt_mean_std(mean, std)
                row[f"{key}_mean"] = mean
                row[f"{key}_std"] = std
            row["Note"] = "strict True/True masks"
        else:
            src = source[source_model]
            for _, key in MAIN_METRICS:
                mean = float(src[f"{key}_mean"])
                std = float(src[f"{key}_std"])
                row[key] = fmt_mean_std(mean, std)
                row[f"{key}_mean"] = mean
                row[f"{key}_std"] = std
            row["Note"] = ""
        rows.append(row)
    return rows


def build_mask_rows() -> list[dict]:
    mapping = {
        "Cold Item R@10": "full_cold_item_macro_r10",
        "Cold Item N@10": "full_cold_item_macro_n10",
        "Cold Item R@20": "full_cold_item_macro_r20",
        "Cold Item N@20": "full_cold_item_macro_n20",
        "Hot Item N@10": "full_hot_item_macro_n10",
        "Interaction Cold N@10": "full_cold_n10",
        "Interaction Hot N@10": "full_hot_n10",
    }
    rows = []
    for name, run_dir in MASK_RUNS:
        manifest = read_manifest(run_dir / MANIFEST)
        cfg = manifest.get("model_config", {})
        values = read_csv_row(run_dir / FINAL_FULLRANK)
        row = {
            "Config": name,
            "MaskKnownPosNeg": str(cfg.get("mask_known_pos_neg")),
            "MaskSameItemNeg": str(cfg.get("mask_same_item_neg")),
        }
        for label, key in mapping.items():
            row[label] = fmt_value(float(values[key]))
            row[f"{label}_raw"] = float(values[key])
        rows.append(row)
    return rows


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, font_size: int = 8) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)


def add_main_table(doc: Document, rows: list[dict]) -> None:
    doc.add_heading("Table 1. Junyi Main Results (3 Seeds)", level=1)
    note = doc.add_paragraph()
    note.add_run("Values are mean +/- std over seeds 2025/2026/2027. ")
    note.add_run("Primary metrics are Cold ItemMacro full-ranking metrics. ")
    note.add_run("Ours is recomputed with strict True/True masks.").italic = True

    headers = ["Model"] + [label for label, _ in MAIN_METRICS]
    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, font_size=8)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")

    best_by_metric = {}
    for _, key in MAIN_METRICS:
        best_by_metric[key] = max(row[f"{key}_mean"] for row in rows)

    for source in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], source["Model"], bold=(source["Model"] == "Ours"))
        for idx, (_, key) in enumerate(MAIN_METRICS, start=1):
            is_best = math.isclose(source[f"{key}_mean"], best_by_metric[key], rel_tol=1e-12, abs_tol=1e-12)
            set_cell_text(cells[idx], source[key], bold=is_best or source["Model"] == "Ours")


def add_mask_table(doc: Document, rows: list[dict]) -> None:
    doc.add_heading("Table 2. Junyi Mask Ablation (seed2025)", level=1)
    note = doc.add_paragraph()
    note.add_run("Single-seed diagnostic table. ")
    note.add_run("Use it as ablation evidence, not as the final multi-seed main result.").italic = True

    headers = [
        "Config",
        "MaskKnown",
        "MaskSame",
        "Cold Item R@10",
        "Cold Item N@10",
        "Cold Item N@20",
        "Hot Item N@10",
        "Interaction Cold N@10",
        "Interaction Hot N@10",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, font_size=8)
        set_cell_shading(table.rows[0].cells[idx], "E2F0D9")

    best_cold_n10 = max(row["Cold Item N@10_raw"] for row in rows)
    for row in rows:
        cells = table.add_row().cells
        values = [
            row["Config"],
            row["MaskKnownPosNeg"],
            row["MaskSameItemNeg"],
            row["Cold Item R@10"],
            row["Cold Item N@10"],
            row["Cold Item N@20"],
            row["Hot Item N@10"],
            row["Interaction Cold N@10"],
            row["Interaction Hot N@10"],
        ]
        for idx, value in enumerate(values):
            bold = row["Config"] == "True/True" or (
                headers[idx] == "Cold Item N@10"
                and math.isclose(row["Cold Item N@10_raw"], best_cold_n10, rel_tol=1e-12, abs_tol=1e-12)
            )
            set_cell_text(cells[idx], value, bold=bold)


def write_csv(rows: list[dict], path: Path, keys: list[str]) -> None:
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=keys)
        writer.writeheader()
        for row in rows:
            writer.writerow({key: row.get(key, "") for key in keys})


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    main_rows = build_main_rows()
    mask_rows = build_mask_rows()

    main_keys = ["Model", "n", "Seeds"] + [key for _, key in MAIN_METRICS] + ["Note"]
    write_csv(main_rows, CSV_PATH, main_keys)

    mask_keys = [
        "Config",
        "MaskKnownPosNeg",
        "MaskSameItemNeg",
        "Cold Item R@10",
        "Cold Item N@10",
        "Cold Item R@20",
        "Cold Item N@20",
        "Hot Item N@10",
        "Interaction Cold N@10",
        "Interaction Hot N@10",
    ]
    write_csv(mask_rows, MASK_CSV_PATH, mask_keys)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    styles["Normal"].font.size = Pt(9)

    title = doc.add_heading("Junyi Dataset Results Tables", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_main_table(doc, main_rows)
    doc.add_paragraph()
    add_mask_table(doc, mask_rows)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.add_run("Source files: ").bold = True
    foot.add_run(str(MAIN_SUMMARY))
    foot.add_run("; strict True/True Ours uses seed2025 from mask_ablation/mask_tt and seed2026/2027 from main_table_3seed.")

    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {CSV_PATH}")
    print(f"Wrote {MASK_CSV_PATH}")


if __name__ == "__main__":
    main()
