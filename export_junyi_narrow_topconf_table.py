import csv
import json
import math
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"D:\DeskTop\MOOCCube")
TEMPLATE = ROOT / "output" / "doc" / "final_narrow_topconf" / "tables_narrow_topconf.docx"
OUT_DIR = ROOT / "output" / "doc" / "junyi_results"

COLD_DOCX_OUT = OUT_DIR / "junyi_main_table_narrow_topconf.docx"
COLD_CSV_OUT = OUT_DIR / "junyi_main_table_narrow_topconf.csv"
HOT_DOCX_OUT = OUT_DIR / "junyi_hot_table_narrow_topconf.docx"
HOT_CSV_OUT = OUT_DIR / "junyi_hot_table_narrow_topconf.csv"
BOTH_DOCX_OUT = OUT_DIR / "junyi_tables_narrow_topconf_cold_hot.docx"

SEED_DIRS = {
    2025: ROOT / "outputs" / "junyi" / "official_prereq_seed2025" / "strict_item_cold_balanced_thr1_seed_2025",
    2026: ROOT / "outputs" / "junyi" / "main_table_3seed" / "strict_item_cold_balanced_thr1_seed_2026",
    2027: ROOT / "outputs" / "junyi" / "main_table_3seed" / "strict_item_cold_balanced_thr1_seed_2027",
}

OURS_TT_DIRS = {
    2025: ROOT / "outputs" / "junyi" / "mask_ablation" / "mask_tt" / "strict_item_cold_balanced_thr1_seed_2025",
    2026: ROOT / "outputs" / "junyi" / "main_table_3seed" / "strict_item_cold_balanced_thr1_seed_2026",
    2027: ROOT / "outputs" / "junyi" / "main_table_3seed" / "strict_item_cold_balanced_thr1_seed_2027",
}

CGRC_STRICTFIX_DIRS = {
    2025: ROOT / "outputs" / "junyi" / "mask_ablation" / "mask_tt" / "strict_item_cold_balanced_thr1_seed_2025",
    2026: ROOT / "outputs" / "junyi" / "main_table_3seed" / "strict_item_cold_balanced_thr1_seed_2026",
    2027: ROOT / "outputs" / "junyi" / "main_table_3seed" / "strict_item_cold_balanced_thr1_seed_2027",
}

BASELINE_SPECS = [
    ("Popularity", "popularity_compare", "popularity_static_result.json"),
    ("BPR", "bpr_compare", "bpr_static_result.json"),
    ("LightGCN", "lightgcn_compare", "lightgcn_static_result.json"),
    ("DropoutNet", "dropoutnet_compare", "dropoutnet_official_static_result.json"),
    ("Content-CBF", "content_profile_compare", "content_profile_static_result.json"),
    ("CCFCRec", "ccfcrec_compare", "ccfcrec_static_result.json"),
    ("ALDI", "aldi_compare", "aldi_static_result.json"),
    ("CGRC", "cgrc_paper_compare_strictfix", "cgrc_paper_static_result.json"),
]

KS = ["5", "10", "20"]
KINDS = ["R", "N"]


def make_metrics(split: str) -> list[tuple[str, str]]:
    return [(f"{split} {kind}@{k}", f"{kind}@{k}") for kind in KINDS for k in KS]


COLD_METRICS = make_metrics("Cold")
HOT_METRICS = make_metrics("Hot")


def read_json_result(path: Path) -> dict:
    data = json.loads(path.read_text(encoding="utf-8"))
    if isinstance(data, dict) and "value" in data:
        return data["value"][0]
    if isinstance(data, list):
        return data[0]
    return data


def read_csv_row(path: Path) -> dict:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return next(csv.DictReader(f))


def mean(values: list[float]) -> float:
    return sum(values) / len(values)


def fmt(value: float) -> str:
    return f"{value:.4f}"


def fmt_imp(value: float) -> str:
    return f"{value:+.2f}%"


def load_baseline(model: str, subdir: str, filename: str, macro_key: str, metrics: list[tuple[str, str]]) -> dict:
    per_metric = {key: [] for _, key in metrics}
    seed_dirs = CGRC_STRICTFIX_DIRS if subdir == "cgrc_paper_compare_strictfix" else SEED_DIRS
    for seed, seed_dir in seed_dirs.items():
        path = seed_dir / subdir / filename
        if not path.exists():
            raise FileNotFoundError(f"Missing {model} seed={seed}: {path}")
        row = read_json_result(path)
        macro = row.get(macro_key)
        if not isinstance(macro, dict):
            raise ValueError(f"Missing {macro_key} in {path}")
        for _, key in metrics:
            per_metric[key].append(float(macro[key]))
    return {key: mean(values) for key, values in per_metric.items()}


def load_ours_tt(csv_prefix: str, metrics: list[tuple[str, str]]) -> dict:
    per_metric = {key: [] for _, key in metrics}
    suffix = {
        "R@5": "r5",
        "R@10": "r10",
        "R@20": "r20",
        "N@5": "n5",
        "N@10": "n10",
        "N@20": "n20",
    }
    for seed, run_dir in OURS_TT_DIRS.items():
        manifest = json.loads((run_dir / "static_protocol_manifest.json").read_text(encoding="utf-8"))
        cfg = manifest.get("model_config", {})
        if cfg.get("mask_known_pos_neg") is not True or cfg.get("mask_same_item_neg") is not True:
            raise ValueError(f"Ours seed={seed} is not strict True/True: {run_dir}")
        row = read_csv_row(run_dir / "final_fullrank_usim_feedback_fast3_content_delta_static.csv")
        for _, key in metrics:
            per_metric[key].append(float(row[f"{csv_prefix}_{suffix[key]}"]))
    return {key: mean(values) for key, values in per_metric.items()}


def build_rows(split: str) -> tuple[list[dict], list[tuple[str, str]]]:
    if split == "Cold":
        metrics = COLD_METRICS
        macro_key = "full_cold_item_macro"
        csv_prefix = "full_cold_item_macro"
    elif split == "Hot":
        metrics = HOT_METRICS
        macro_key = "full_hot_item_macro"
        csv_prefix = "full_hot_item_macro"
    else:
        raise ValueError(split)

    rows = []
    for model, subdir, filename in BASELINE_SPECS:
        values = load_baseline(model, subdir, filename, macro_key, metrics)
        rows.append({"Model": model, **values})
    rows.append({"Model": "Ours", **load_ours_tt(csv_prefix, metrics)})

    imp = {"Model": "Imp."}
    for _, key in metrics:
        ours = rows[-1][key]
        best_baseline = max(row[key] for row in rows[:-1])
        imp[key] = ((ours / best_baseline) - 1.0) * 100.0 if best_baseline else 0.0
    rows.append(imp)
    return rows, metrics


def write_csv(rows: list[dict], metrics: list[tuple[str, str]], path: Path) -> None:
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Model"] + [label for label, _ in metrics])
        for row in rows:
            if row["Model"] == "Imp.":
                writer.writerow([row["Model"]] + [fmt_imp(row[key]) for _, key in metrics])
            else:
                writer.writerow([row["Model"]] + [fmt(row[key]) for _, key in metrics])


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_paragraph_text(paragraph, text: str) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(10)


def set_cell_text(cell, text: str, *, bold: bool = False, underline: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(9)
    run.bold = bold
    run.underline = underline


def ensure_table_rows(table, n_rows: int) -> None:
    while len(table.rows) < n_rows:
        table.add_row()


def fill_table(table, rows: list[dict], metrics: list[tuple[str, str]]) -> None:
    ensure_table_rows(table, len(rows) + 1)
    headers = ["Model"] + [label for label, _ in metrics]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)

    data_rows = rows[:-1]
    best = {key: max(row[key] for row in data_rows) for _, key in metrics}
    second = {}
    for _, key in metrics:
        vals = sorted({row[key] for row in data_rows}, reverse=True)
        second[key] = vals[1] if len(vals) > 1 else vals[0]

    for row_idx, row in enumerate(rows, start=1):
        cells = table.rows[row_idx].cells
        set_cell_text(cells[0], row["Model"], bold=(row["Model"] == "Ours"))
        for col_idx, (_, key) in enumerate(metrics, start=1):
            if row["Model"] == "Imp.":
                set_cell_text(cells[col_idx], fmt_imp(row[key]))
                continue
            value = row[key]
            is_best = math.isclose(value, best[key], rel_tol=1e-12, abs_tol=1e-12)
            is_second = math.isclose(value, second[key], rel_tol=1e-12, abs_tol=1e-12)
            set_cell_text(cells[col_idx], fmt(value), bold=is_best, underline=(is_second and not is_best))


def write_single_docx(
    split: str,
    rows: list[dict],
    metrics: list[tuple[str, str]],
    out_path: Path,
) -> None:
    shutil.copyfile(TEMPLATE, out_path)
    doc = Document(out_path)
    if len(doc.tables) > 1:
        remove_element(doc.tables[1]._element)
    for paragraph in list(doc.paragraphs):
        if paragraph.text.startswith("Table 5:") or paragraph.text.startswith("All values are three-seed means. The main method"):
            remove_element(paragraph._element)

    set_paragraph_text(
        doc.paragraphs[0],
        f"Table: {split}-item performance comparison on the Junyi balanced static item-cold split. The best result is in bold and the second-best result is underlined.",
    )
    set_paragraph_text(
        doc.paragraphs[1],
        "All values are three-seed means under full-ranking item-macro evaluation. Ours uses strict True/True masks.",
    )
    fill_table(doc.tables[0], rows, metrics)
    doc.save(out_path)


def write_both_docx(cold_rows, cold_metrics, hot_rows, hot_metrics) -> None:
    shutil.copyfile(TEMPLATE, BOTH_DOCX_OUT)
    doc = Document(BOTH_DOCX_OUT)
    set_paragraph_text(
        doc.paragraphs[0],
        "Table: Cold-item performance comparison on the Junyi balanced static item-cold split. The best result is in bold and the second-best result is underlined.",
    )
    set_paragraph_text(
        doc.paragraphs[1],
        "All values are three-seed means under full-ranking item-macro evaluation. Ours uses strict True/True masks.",
    )
    set_paragraph_text(
        doc.paragraphs[2],
        "Table: Hot-item performance comparison on the Junyi balanced static item-cold split. The best result is in bold and the second-best result is underlined.",
    )
    set_paragraph_text(
        doc.paragraphs[3],
        "All values are three-seed means under full-ranking item-macro evaluation. Ours uses strict True/True masks.",
    )
    fill_table(doc.tables[0], cold_rows, cold_metrics)
    fill_table(doc.tables[1], hot_rows, hot_metrics)
    doc.save(BOTH_DOCX_OUT)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    cold_rows, cold_metrics = build_rows("Cold")
    hot_rows, hot_metrics = build_rows("Hot")

    write_csv(cold_rows, cold_metrics, COLD_CSV_OUT)
    write_csv(hot_rows, hot_metrics, HOT_CSV_OUT)
    write_single_docx("Cold", cold_rows, cold_metrics, COLD_DOCX_OUT)
    write_single_docx("Hot", hot_rows, hot_metrics, HOT_DOCX_OUT)
    write_both_docx(cold_rows, cold_metrics, hot_rows, hot_metrics)

    print(f"Wrote {COLD_DOCX_OUT}")
    print(f"Wrote {COLD_CSV_OUT}")
    print(f"Wrote {HOT_DOCX_OUT}")
    print(f"Wrote {HOT_CSV_OUT}")
    print(f"Wrote {BOTH_DOCX_OUT}")


if __name__ == "__main__":
    main()
