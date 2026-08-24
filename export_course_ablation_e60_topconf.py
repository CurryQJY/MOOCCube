"""Export the E60 course-component ablation table in compact top-conf style.

Inputs:
    outputs/content_delta_pop5/course_ablation_e60_3seed_corrected/*/fast3_static_multiseed_summary.csv
    (override with USIM_COURSE_ABLATION_ROOT)

Outputs:
    output/doc/course_ablation_e60_3seed_topconf/ablation_table_narrow_topconf.docx
    output/doc/course_ablation_e60_3seed_topconf/ablation_table_narrow_topconf.csv
    output/doc/course_ablation_e60_3seed_topconf/ablation_table_narrow_topconf_mean_std.csv
    output/doc/course_ablation_e60_3seed_topconf/ablation_table_narrow_topconf.tex
"""

from __future__ import annotations

import os
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(
    os.environ.get(
        "USIM_COURSE_ABLATION_ROOT",
        "outputs/content_delta_pop5/course_ablation_e60_3seed_corrected",
    )
)
OUT_DIR = Path("output/doc/course_ablation_e60_3seed_topconf")

METRICS = [
    ("Cold R@5", "full_cold_item_macro_r5"),
    ("Cold R@10", "full_cold_item_macro_r10"),
    ("Cold R@20", "full_cold_item_macro_r20"),
    ("Cold N@5", "full_cold_item_macro_n5"),
    ("Cold N@10", "full_cold_item_macro_n10"),
    ("Cold N@20", "full_cold_item_macro_n20"),
]

ROWS = [
    ("Ours", ROOT / "full" / "fast3_static_multiseed_summary.csv"),
    ("w/o Course-aware Reward", ROOT / "wo_course_reward" / "fast3_static_multiseed_summary.csv"),
    ("w/o Course-aware User Selection", ROOT / "wo_course_candidate" / "fast3_static_multiseed_summary.csv"),
    ("w/o Prereq. Auxiliary Loss", ROOT / "wo_prereq_aux" / "fast3_static_multiseed_summary.csv"),
    ("w/o All Course Signals", ROOT / "wo_all_course_signals" / "fast3_static_multiseed_summary.csv"),
]


def metric_mean(row: pd.Series, key: str) -> float:
    return float(row[f"{key}_mean"])


def metric_std(row: pd.Series, key: str) -> float:
    return float(row[f"{key}_std"])


def load_table() -> pd.DataFrame:
    rows = []
    for name, path in ROWS:
        if not path.exists():
            raise FileNotFoundError(f"Missing ablation summary: {path}")
        raw = pd.read_csv(path)
        if raw.empty:
            raise ValueError(f"Empty ablation summary: {path}")
        src = raw.iloc[0]
        out = {
            "Model": name,
            "runs": src.get("runs", ""),
            "seeds": src.get("seeds", ""),
        }
        for label, key in METRICS:
            out[label] = metric_mean(src, key)
            out[f"{label} std"] = metric_std(src, key)
        rows.append(out)
    return pd.DataFrame(rows)


def fmt(value: float) -> str:
    return f"{float(value):.4f}"


def fmt_mean_std(mean: float, std: float) -> str:
    return f"{float(mean):.4f} (+/-{float(std):.4f})"


def best_second(df: pd.DataFrame) -> dict[str, tuple[float | None, float | None]]:
    out = {}
    for label, _ in METRICS:
        values = sorted({float(v) for v in df[label].dropna().tolist()}, reverse=True)
        out[label] = (values[0] if values else None, values[1] if len(values) > 1 else None)
    return out


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, spec in kwargs.items():
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in spec.items():
            element.set(qn(f"w:{key}"), str(value))


def clear_cell_borders(cell) -> None:
    nil = {"val": "nil", "sz": "0", "space": "0", "color": "FFFFFF"}
    set_cell_border(cell, top=nil, bottom=nil, left=nil, right=nil, insideH=nil, insideV=nil)


def add_rule(cell, edge: str, width: str = "6") -> None:
    set_cell_border(cell, **{edge: {"val": "single", "sz": width, "space": "0", "color": "000000"}})


def set_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=28, start=24, bottom=28, end=24) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_text(cell, text: str, *, bold=False, underline=False, italic=False, size=8.7, align=None) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.underline = underline
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def style_cell(cell) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    clear_cell_borders(cell)


def apply_rules(table, separator_rows: tuple[int, ...] = ()) -> None:
    n_rows = len(table.rows)
    n_cols = len(table.columns)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            clear_cell_borders(cell)
            if col_idx < n_cols - 1:
                add_rule(cell, "right", "5")
            if row_idx == 0:
                add_rule(cell, "top", "9")
                add_rule(cell, "bottom", "7")
            if row_idx in separator_rows:
                add_rule(cell, "bottom", "6")
            if row_idx == n_rows - 1:
                add_rule(cell, "bottom", "9")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(9.2)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(7.5)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def build_docx(df: pd.DataFrame, out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(9)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_caption(
        doc,
        (
            "Table 2: Cold-item ablation study on the balanced static item-cold split. "
            "The best result is in bold and the second-best result is underlined."
        ),
    )
    ranks = best_second(df)
    table = doc.add_table(rows=len(df) + 1, cols=1 + len(METRICS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.42] + [0.58] * len(METRICS)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_width(cell, widths[idx])
            style_cell(cell)

    headers = ["Model"] + [label for label, _ in METRICS]
    for idx, header in enumerate(headers):
        add_text(table.rows[0].cells[idx], header, bold=(idx == 0), italic=(idx > 0), size=7.8)

    for row_idx, (_, row) in enumerate(df.iterrows(), start=1):
        model = str(row["Model"])
        add_text(table.rows[row_idx].cells[0], model, bold=True, italic=(model == "Ours"), size=8.2)
        for col_idx, (label, _) in enumerate(METRICS, start=1):
            value = float(row[label])
            best, second = ranks[label]
            is_best = best is not None and abs(value - best) < 5e-13
            is_second = second is not None and abs(value - second) < 5e-13
            add_text(table.rows[row_idx].cells[col_idx], fmt(value), bold=is_best, underline=is_second, size=8.0)

    apply_rules(table, separator_rows=(1,))
    add_note(doc, "All values are three-seed means under full-ranking item-macro evaluation.")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def latex_table(df: pd.DataFrame) -> str:
    ranks = best_second(df)
    lines = [
        "\\begin{table}[t]",
        "\\centering",
        "\\small",
        "\\caption{Cold-item ablation study on the balanced static item-cold split.}",
        "\\label{tab:course-ablation}",
        "\\begin{tabular}{l" + "c" * len(METRICS) + "}",
        "\\toprule",
        "Model & " + " & ".join(label for label, _ in METRICS) + " \\\\",
        "\\midrule",
    ]
    for idx, row in df.iterrows():
        if idx == 1:
            lines.append("\\midrule")
        vals = []
        for label, _ in METRICS:
            value = float(row[label])
            best, second = ranks[label]
            text = fmt(value)
            if best is not None and abs(value - best) < 5e-13:
                text = f"\\textbf{{{text}}}"
            elif second is not None and abs(value - second) < 5e-13:
                text = f"\\underline{{{text}}}"
            vals.append(text)
        lines.append(f"{row['Model']} & " + " & ".join(vals) + " \\\\")
    lines.extend(["\\bottomrule", "\\end{tabular}", "\\end{table}", ""])
    return "\n".join(lines)


def export_csvs(df: pd.DataFrame) -> None:
    paper_rows = []
    for _, row in df.iterrows():
        out = {"Model": row["Model"]}
        for label, _ in METRICS:
            out[label] = fmt(row[label])
        paper_rows.append(out)
    pd.DataFrame(paper_rows).to_csv(OUT_DIR / "ablation_table_narrow_topconf.csv", index=False)

    rows = []
    for _, row in df.iterrows():
        out = {"Model": row["Model"], "runs": row["runs"], "seeds": row["seeds"]}
        for label, _ in METRICS:
            out[label] = fmt_mean_std(row[label], row[f"{label} std"])
        rows.append(out)
    pd.DataFrame(rows).to_csv(OUT_DIR / "ablation_table_narrow_topconf_mean_std.csv", index=False)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    df = load_table()
    export_csvs(df)
    (OUT_DIR / "ablation_table_narrow_topconf.tex").write_text(latex_table(df), encoding="utf-8")
    build_docx(df, OUT_DIR / "ablation_table_narrow_topconf.docx")
    print(f"Wrote {OUT_DIR / 'ablation_table_narrow_topconf.csv'}")
    print(f"Wrote {OUT_DIR / 'ablation_table_narrow_topconf_mean_std.csv'}")
    print(f"Wrote {OUT_DIR / 'ablation_table_narrow_topconf.tex'}")
    print(f"Wrote {OUT_DIR / 'ablation_table_narrow_topconf.docx'}")


if __name__ == "__main__":
    main()
