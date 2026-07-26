"""Export the narrow static item-cold main table to a Word document."""

import os
from pathlib import Path
import shutil

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("outputs/content_delta_pop5/static_item_cold/strict_item_cold_thr1_seed_2025")
IN_CSV = Path(os.environ.get("MAIN_TABLE_CSV", str(ROOT / "main_table_fair_v1" / "main_table_fair_v1_paper_narrow.csv")))
OUT_DOCX = Path(os.environ.get("MAIN_TABLE_DOCX", "output/doc/main_table_fair_v1_paper_narrow_topconf.docx"))
KDD_DOCX = Path(os.environ.get("MAIN_TABLE_KDD_DOCX", "output/doc/main_table_fair_v1_paper_narrow_kdd_style.docx"))
REF_DOCX = Path(os.environ.get("MAIN_TABLE_REF_DOCX", "output/doc/main_table_fair_v1_paper_reference_structure_v3.docx"))
COMPAT_DOCX = Path(os.environ.get("MAIN_TABLE_COMPAT_DOCX", "output/doc/main_table_fair_v1_paper_narrow.docx"))

METHOD_GROUPS = [
    ("Heuristic", ["Popularity"]),
    ("ID-based CF", ["BPR", "LightGCN"]),
    ("Cold-start baselines", ["DropoutNet", "GAR", "ContentProfile", "CGRC", "ALDI"]),
    ("Ours", ["FAST3"]),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn("w:{}".format(key)), str(value))


def clear_cell_borders(cell) -> None:
    nil = {"val": "nil", "sz": "0", "space": "0", "color": "FFFFFF"}
    set_cell_border(
        cell,
        top=nil,
        bottom=nil,
        left=nil,
        right=nil,
        insideH=nil,
        insideV=nil,
    )


def add_rule(cell, edge: str, width: str = "8", color: str = "000000") -> None:
    set_cell_border(cell, **{edge: {"val": "single", "sz": width, "space": "0", "color": color}})


def set_row_height(row, height_twips: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(height_twips))
    tr_height.set(qn("w:hRule"), "atLeast")


def set_cell_margins(cell, top=60, start=60, bottom=60, end=60) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:{}".format(margin)))
        if node is None:
            node = OxmlElement("w:{}".format(margin))
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_text(cell, text: str, bold=False, underline=False, italic=False, size=8.5, align=None) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.underline = underline
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_text_to_new_paragraph(cell, text: str, bold=False, italic=False, size=9.0) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_caption_paragraph(doc: Document, text: str, size: float = 8.0) -> None:
    caption = doc.add_paragraph()
    caption.style = doc.styles["Normal"]
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.space_after = Pt(2)
    caption.paragraph_format.space_before = Pt(0)
    run = caption.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def format_value(value) -> str:
    return "{:.4f}".format(float(value))


def format_imp(value: float) -> str:
    sign = "+" if value >= 0 else ""
    return "{}{:.2f}%".format(sign, value * 100.0)


def best_and_second(df: pd.DataFrame, cols) -> dict:
    out = {}
    for col in cols:
        vals = sorted({float(v) for v in df[col].dropna().tolist()}, reverse=True)
        out[col] = (vals[0] if vals else None, vals[1] if len(vals) > 1 else None)
    return out


def style_table_cell(cell) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=35, start=25, bottom=35, end=25)
    clear_cell_borders(cell)


def build_docx(df: pd.DataFrame, out_path: Path) -> None:
    metric_cols = ["Cold R@10", "Cold N@10", "Hot R@10", "Hot N@10"]
    ranks = best_and_second(df, metric_cols)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.start_type = WD_SECTION.CONTINUOUS

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(9)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(2)
    caption.paragraph_format.space_before = Pt(0)
    add_caption = caption.add_run(
        "Table 1: Static item-cold recommendation results under full-ranking evaluation."
    )
    add_caption.bold = True
    add_caption.font.name = "Times New Roman"
    add_caption.font.size = Pt(8)
    add_caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    table = doc.add_table(rows=len(df) + 2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [1.25, 0.55, 0.55, 0.55, 0.55]
    for row in table.rows:
        for idx, width in enumerate(widths):
            set_width(row.cells[idx], width)

    hdr0 = table.rows[0].cells
    hdr0[0].merge(table.rows[1].cells[0])
    hdr0[1].merge(hdr0[2])
    hdr0[3].merge(hdr0[4])
    add_text(table.rows[0].cells[0], "Model", bold=True, size=8)
    add_text(table.rows[0].cells[1], "Cold", bold=True, size=8)
    add_text(table.rows[0].cells[3], "Hot", bold=True, size=8)

    hdr1 = table.rows[1].cells
    add_text(hdr1[1], "R@10", bold=True, size=8)
    add_text(hdr1[2], "N@10", bold=True, size=8)
    add_text(hdr1[3], "R@10", bold=True, size=8)
    add_text(hdr1[4], "N@10", bold=True, size=8)

    for row in table.rows[:2]:
        set_repeat_table_header(row)
        for cell in row.cells:
            style_table_cell(cell)

    # Booktabs-style rules: no vertical grid, no shading.
    for cell in table.rows[0].cells:
        add_rule(cell, "top", width="12")
    for cell in (table.rows[0].cells[1], table.rows[0].cells[3]):
        add_rule(cell, "bottom", width="6")
    for cell in table.rows[1].cells:
        add_rule(cell, "bottom", width="8")

    for row_idx, (_, row) in enumerate(df.iterrows(), start=2):
        cells = table.rows[row_idx].cells
        model_name = str(row["Model"])
        add_text(cells[0], model_name, bold=(model_name == "FAST3"), size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
        style_table_cell(cells[0])

        for col_idx, col in enumerate(metric_cols, start=1):
            value = float(row[col])
            best, second = ranks[col]
            is_best = best is not None and abs(value - best) < 5e-13
            is_second = second is not None and abs(value - second) < 5e-13
            add_text(cells[col_idx], format_value(value), bold=is_best, underline=is_second, size=8)
            style_table_cell(cells[col_idx])

    for cell in table.rows[-1].cells:
        add_rule(cell, "bottom", width="12")

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note.paragraph_format.space_before = Pt(2)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(
        "Note: Cold threshold = 1 training interaction; test history = train-only. "
        "Best results are bolded and second-best results are underlined."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(7)
    run.italic = True
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def build_kdd_docx(df: pd.DataFrame, out_path: Path) -> None:
    metric_cols = ["Cold R@10", "Cold N@10", "Hot R@10", "Hot N@10"]
    ranks = best_and_second(df, metric_cols)
    model_to_row = {str(row["Model"]): row for _, row in df.iterrows()}

    grouped_rows = []
    for group_name, models in METHOD_GROUPS:
        grouped_rows.append(("__group__", group_name, None))
        for model_name in models:
            if model_name in model_to_row:
                grouped_rows.append(("__model__", model_name, model_to_row[model_name]))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    section.start_type = WD_SECTION.CONTINUOUS

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(8)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_caption_paragraph(
        doc,
        "Table 1. Overall performance on the static item-cold split. "
        "Bold and underline denote the best and second-best results.",
        size=8,
    )

    table = doc.add_table(rows=len(grouped_rows) + 2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [1.38, 0.54, 0.54, 0.54, 0.54]
    for row in table.rows:
        set_row_height(row, 230)
        for idx, width in enumerate(widths):
            set_width(row.cells[idx], width)
            style_table_cell(row.cells[idx])

    hdr0 = table.rows[0].cells
    hdr0[0].merge(table.rows[1].cells[0])
    hdr0[1].merge(hdr0[2])
    hdr0[3].merge(hdr0[4])
    add_text(table.rows[0].cells[0], "Method", bold=True, size=7.8, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_text(table.rows[0].cells[1], "Cold", bold=True, size=7.8)
    add_text(table.rows[0].cells[3], "Hot", bold=True, size=7.8)

    hdr1 = table.rows[1].cells
    add_text(hdr1[1], "R@10", bold=True, size=7.8)
    add_text(hdr1[2], "N@10", bold=True, size=7.8)
    add_text(hdr1[3], "R@10", bold=True, size=7.8)
    add_text(hdr1[4], "N@10", bold=True, size=7.8)

    for row in table.rows[:2]:
        set_repeat_table_header(row)
    for cell in table.rows[0].cells:
        add_rule(cell, "top", width="12")
    for cell in (table.rows[0].cells[1], table.rows[0].cells[3]):
        add_rule(cell, "bottom", width="6")
    for cell in table.rows[1].cells:
        add_rule(cell, "bottom", width="8")

    out_row = 2
    for kind, label, row in grouped_rows:
        cells = table.rows[out_row].cells
        if kind == "__group__":
            cells[0].merge(cells[4])
            add_text(cells[0], label, italic=True, size=7.6, align=WD_ALIGN_PARAGRAPH.LEFT)
            for cell in table.rows[out_row].cells:
                style_table_cell(cell)
            add_rule(cells[0], "top", width="4", color="808080")
        else:
            method = label
            display = "FAST3 (ours)" if method == "FAST3" else method
            add_text(cells[0], display, bold=(method == "FAST3"), size=7.8, align=WD_ALIGN_PARAGRAPH.LEFT)
            for col_idx, col in enumerate(metric_cols, start=1):
                value = float(row[col])
                best, second = ranks[col]
                is_best = best is not None and abs(value - best) < 5e-13
                is_second = second is not None and abs(value - second) < 5e-13
                add_text(cells[col_idx], format_value(value), bold=is_best, underline=is_second, size=7.8)
        out_row += 1

    for cell in table.rows[-1].cells:
        add_rule(cell, "bottom", width="12")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def build_reference_structure_docx(df: pd.DataFrame, out_path: Path) -> None:
    metric_cols = ["Cold R@10", "Cold N@10", "Hot R@10", "Hot N@10"]
    ranks = best_and_second(df, metric_cols)

    model_order = [
        "Popularity",
        "BPR",
        "LightGCN",
        "DropoutNet",
        "GAR",
        "ContentProfile",
        "CGRC",
        "ALDI",
        "FAST3",
    ]
    model_to_row = {str(row["Model"]): row for _, row in df.iterrows()}
    rows = [model_to_row[name] for name in model_order if name in model_to_row]
    fast_row = model_to_row["FAST3"]
    baseline_rows = [row for row in rows if str(row["Model"]) != "FAST3"]
    best_baseline = {
        col: max(float(row[col]) for row in baseline_rows)
        for col in metric_cols
    }

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.start_type = WD_SECTION.CONTINUOUS

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.space_before = Pt(0)
    cap = caption.add_run(
        "Table 1: Performance comparison on the static item-cold split with K = 10. "
        "The best performance is indicated in bold, while the second-best performance is underlined. "
        "The relative improvements over the best baseline are denoted as Imp."
    )
    cap.bold = True
    cap.font.name = "Times New Roman"
    cap.font.size = Pt(11)
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    table = doc.add_table(rows=len(rows) + 3, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.62, 1.02, 1.02, 1.02, 1.02]
    for row in table.rows:
        set_row_height(row, 310)
        for idx, width in enumerate(widths):
            set_width(row.cells[idx], width)
            style_table_cell(row.cells[idx])

    hdr0 = table.rows[0].cells
    hdr0[0].merge(table.rows[1].cells[0])
    hdr0[1].merge(hdr0[2])
    hdr0[3].merge(hdr0[4])
    add_text(table.rows[0].cells[0], "Model", bold=True, size=11)
    add_text(table.rows[0].cells[1], "Cold", bold=True, size=11)
    add_text(table.rows[0].cells[3], "Hot", bold=True, size=11)

    hdr1 = table.rows[1].cells
    add_text(hdr1[1], "Recall@10", italic=True, size=10.5)
    add_text(hdr1[2], "NDCG@10", italic=True, size=10.5)
    add_text(hdr1[3], "Recall@10", italic=True, size=10.5)
    add_text(hdr1[4], "NDCG@10", italic=True, size=10.5)

    for row in table.rows[:2]:
        set_repeat_table_header(row)

    # Horizontal rules matching the reference structure.
    out_row = 2
    for row in rows:
        model_name = str(row["Model"])
        cells = table.rows[out_row].cells
        display = "FAST3" if model_name == "FAST3" else model_name
        add_text(
            cells[0],
            display,
            bold=True,
            italic=(model_name == "FAST3"),
            size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        for col_idx, col in enumerate(metric_cols, start=1):
            value = float(row[col])
            best, second = ranks[col]
            is_best = best is not None and abs(value - best) < 5e-13
            is_second = second is not None and abs(value - second) < 5e-13
            add_text(cells[col_idx], format_value(value), bold=is_best, underline=is_second, size=10.5)
        out_row += 1

    imp_cells = table.rows[out_row].cells
    add_text(imp_cells[0], "Imp.", bold=True, size=10.5)
    for col_idx, col in enumerate(metric_cols, start=1):
        imp = (float(fast_row[col]) - best_baseline[col]) / best_baseline[col]
        add_text(imp_cells[col_idx], format_imp(imp), bold=True, size=10.5)

    apply_reference_borders(table)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def apply_reference_borders(table) -> None:
    """Draw continuous borders matching the provided top-conference example."""
    for row in table.rows:
        for cell in row.cells:
            clear_cell_borders(cell)

    n_rows = len(table.rows)
    n_cols = len(table.columns)

    # Top rule across the whole table.
    for cell in table.rows[0].cells:
        add_rule(cell, "top", width="12")

    # Partial cmidrules under the group headers (Cold / Hot), not under Model.
    add_rule(table.rows[0].cells[1], "bottom", width="6")
    add_rule(table.rows[0].cells[3], "bottom", width="6")

    # Header bottom rule across all columns.
    for cell in table.rows[1].cells:
        add_rule(cell, "bottom", width="10")

    # Continuous vertical rules.
    # Header group row: after Model and after Cold.
    add_rule(table.rows[0].cells[0], "right", width="6")
    add_rule(table.rows[0].cells[1], "right", width="6")
    # Metric header and body: after every column except the last one.
    for row_idx in range(1, n_rows):
        row = table.rows[row_idx]
        for col_idx in range(n_cols - 1):
            add_rule(row.cells[col_idx], "right", width="6")

    # Block separators: after LightGCN and after ALDI, like the reference table.
    for row_idx in (4, 9):
        for cell in table.rows[row_idx].cells:
            add_rule(cell, "bottom", width="8")

    # Bottom rule under Imp.
    for cell in table.rows[-1].cells:
        add_rule(cell, "bottom", width="12")


def main() -> None:
    if not IN_CSV.exists():
        raise FileNotFoundError("Missing narrow table CSV: {}".format(IN_CSV))
    df = pd.read_csv(IN_CSV)
    build_docx(df, OUT_DOCX)
    build_kdd_docx(df, KDD_DOCX)
    build_reference_structure_docx(df, REF_DOCX)
    COMPAT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    try:
        shutil.copyfile(REF_DOCX, COMPAT_DOCX)
    except PermissionError:
        print("Skip overwrite locked {}".format(COMPAT_DOCX))
    print("Wrote {}".format(OUT_DOCX))
    print("Wrote {}".format(KDD_DOCX))
    print("Wrote {}".format(REF_DOCX))
    print("Wrote {}".format(COMPAT_DOCX))


if __name__ == "__main__":
    main()
