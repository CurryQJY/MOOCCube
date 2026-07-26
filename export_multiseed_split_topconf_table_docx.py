"""Export multi-seed split statistics in top-conference three-line tables."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


IN_DIR = Path("output/doc/dataset_split_statistics")
OUT_DIR = Path("output/doc/dataset_topconf_table")
OUT_PATH = OUT_DIR / "multiseed_split_statistics_topconf.docx"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def fmt_pct(value: str) -> str:
    text = str(value).strip()
    return text if text.endswith("%") else f"{text}%"


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, spec in kwargs.items():
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in spec.items():
            element.set(qn(f"w:{key}"), str(value))


def clear_cell_borders(cell) -> None:
    nil = {"val": "nil", "sz": "0", "space": "0", "color": "FFFFFF"}
    set_cell_border(cell, top=nil, bottom=nil, left=nil, right=nil)


def add_rule(cell, edge: str, width: str) -> None:
    set_cell_border(cell, **{edge: {"val": "single", "sz": width, "space": "0", "color": "000000"}})


def set_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=35, start=22, bottom=35, end=22) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_text(cell, text: str, *, bold=False, size=9.2) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def apply_three_line_rules(table) -> None:
    n_rows = len(table.rows)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            clear_cell_borders(cell)
            if row_idx == 0:
                add_rule(cell, "top", "10")
                add_rule(cell, "bottom", "7")
            if row_idx == n_rows - 1:
                add_rule(cell, "bottom", "10")


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_three_line_table(
    doc: Document,
    caption: str,
    columns: list[str],
    rows: list[list[str]],
    widths: list[float],
    font_size: float = 9.2,
) -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=len(rows) + 1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_width(cell, widths[idx])
            set_cell_margins(cell)
            clear_cell_borders(cell)

    for idx, col in enumerate(columns):
        add_text(table.rows[0].cells[idx], col, bold=True, size=font_size)
    for r_idx, values in enumerate(rows, start=1):
        for c_idx, value in enumerate(values):
            add_text(table.rows[r_idx].cells[c_idx], value, bold=(c_idx == 0), size=font_size)

    apply_three_line_rules(table)


def add_gap(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)


def build_docx() -> None:
    split_rows = read_csv(IN_DIR / "split_statistics_by_seed.csv")
    cold_rows = read_csv(IN_DIR / "cold_split_statistics_by_seed.csv")

    split_cols = ["Seed", "Train", "Valid", "Test", "Train %", "Valid %", "Test %"]
    split_values = [
        [
            row["Seed"],
            row["Train"],
            row["Val"],
            row["Test"],
            fmt_pct(row["Train %"]),
            fmt_pct(row["Val %"]),
            fmt_pct(row["Test %"]),
        ]
        for row in split_rows
    ]

    cold_cols = ["Seed", "Train Items", "Valid Items", "Test Items", "Valid Cold", "Test Cold"]
    cold_values = [
        [
            row["Seed"],
            row["Train Items"],
            row["Val Items"],
            row["Test Items"],
            row["Val Cold Items"],
            row["Test Cold Items"],
        ]
        for row in cold_rows
    ]

    cold_inter_cols = ["Seed", "Valid Cold Int.", "Test Cold Int.", "Valid SeenU", "Test SeenU"]
    cold_inter_values = [
        [
            split["Seed"],
            cold["Val Cold Int."],
            cold["Test Cold Int."],
            fmt_pct(split["Val SeenU"]),
            fmt_pct(split["Test SeenU"]),
        ]
        for split, cold in zip(split_rows, cold_rows)
    ]

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    add_three_line_table(
        doc,
        "Table 2: The statistics of data splits under three seeds.",
        split_cols,
        split_values,
        widths=[0.68, 1.05, 1.0, 1.0, 0.78, 0.78, 0.78],
        font_size=9.2,
    )
    add_gap(doc)
    add_three_line_table(
        doc,
        "Table 3: The statistics of item splits under three seeds.",
        cold_cols,
        cold_values,
        widths=[0.68, 1.05, 1.05, 1.05, 1.0, 1.0],
        font_size=9.2,
    )
    add_gap(doc)
    add_three_line_table(
        doc,
        "Table 4: The statistics of cold interactions under three seeds.",
        cold_inter_cols,
        cold_inter_values,
        widths=[0.75, 1.45, 1.45, 1.05, 1.05],
        font_size=9.2,
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)


def main() -> None:
    build_docx()
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
